#!/usr/bin/env python3

import sys
import os
import argparse
import re

from lxml import etree 
from docx import Document

if __name__ == '__main__':

   argparser = argparse.ArgumentParser()

   usage = 'Usage : %prog [options]'
   argparser.add_argument('-v', '--specvals', dest = 'vals', nargs = '*',
      help = 'CSV File with Substitution Values')
   argparser.add_argument('-s', '--spec', dest = 'spec', nargs = '*',
      help = 'Word Document to Update (.docx)')
   argparser.add_argument('-d', '--debug', dest = 'debug', action = 'store_true',
      help = 'Turn on debug output in case parsing is failing')
       
   argparser.set_defaults(vals=None, spec=None)
   (options) = argparser.parse_args()

   if options.vals is None or options.spec is None:
      print("Please provide both the CSV file and Word document to update...\n")
      argparser.print_help()
      sys.exit()

   if os.access(*options.vals, os.R_OK):
      vals = open(*options.vals, "r")
   else:
      print("ERROR: Cannot open %s for reading...\n" % options.vals)
      sys.exit()

   if os.access(*options.spec, os.R_OK):
      spec = Document(*options.spec)
   else:
      print("ERROR: Cannot open %s for reading...\n" % options.spec)
      sys.exit()

   # Extract all the key-value pairs from the CSV file
   # Expect to find a header with "Key" and "Value" columns. 
   # Keys are in the form @A000 --> @Z999 
   kv_csv = {}   # This will hold all the KV pairs found in the file
   k_offset = 0
   v_offset = 0
   for l in vals:
      l = l.strip()
      csv = l.split(',')
      if csv:
         if (k_offset == 0 and v_offset == 0):
            for h in csv:
               hs = h.strip()
               if hs == "Key":
                  k_offset = csv.index(h)
                  if options.debug:
                     print("Found key column at offset %d" % k_offset)
               if hs == "Value":
                  v_offset = csv.index(h)
                  if options.debug:
                     print("Found value column at offset %d" % v_offset)
         else: 
            if csv:
               m = re.match(r'@[A-Z][0-9][0-9][0-9]', csv[k_offset].strip())
               if m:
                  key = m.group(0)
                  kv_csv[key] = csv[v_offset].strip()
   vals.close()

   # Extract all the keys from the specification 
   k_spec = []   # This will hold all the keys in the specification
   k_spec_nontable = []   # This will hold all any keys incorrectly placed outside a table
   for t in spec.tables:
      for r in t.rows:
         for c in r.cells:
            for p in c.paragraphs:
               k = re.findall(r'@[A-Z][0-9][0-9][0-9]', p.text)
               for i in k:
                  k_spec.append(i)
                  if options.debug:
                     print("Found key in Word document: %s" % i)
   for p in spec.paragraphs:
      k = re.findall(r'@[A-Z][0-9][0-9][0-9]', p.text)
      for i in k:
         k_spec_nontable.append(i)
   if k_spec_nontable:
         print("WARNING: The following keys in the document are not contained in a table\n")
         print(k_spec_nontable)

   # Make sure that all the keys found in the specification are entered in the CSV
   k_missing = []   # This will hold any keys that cannot be found in the CSV file 
   for key in k_spec:
      if not key in kv_csv:
         k_missing.append(key)
   if k_missing:
         print("ERROR: The following keys in the specification are missing from the CSV file:\n")
         print(k_missing)
         sys.exit()

   # Make substitutions in the document
   for t in spec.tables:
      for r in t.rows:
         for c in r.cells:
            for p in c.paragraphs:
               k = re.findall(r'@[A-Z][0-9][0-9][0-9]', p.text)
               for i in k:
                  if options.debug:
                     print('Replacing %s with %s' % (i, kv_csv[i]))
                  p.text = p.text.replace(i, kv_csv[i])
                  if 'Table Contents' in spec.styles:
                     p.style = 'Table Contents'
      
   spec.save(str(*options.spec) + '.populated.docx')
   sys.exit()
